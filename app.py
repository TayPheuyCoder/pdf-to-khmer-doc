import streamlit as st
import pdfplumber
from docx import Document
import pytesseract
from pdf2image import convert_from_bytes
import io
import re
import openai
import os

from deep_translator import GoogleTranslator
from langdetect import detect

# ---------------------------------
# CONFIG
# ---------------------------------
OCR_DPI = 200
TRANSLATE_LIMIT = 4500

# ---------------------------------
# OPENAI API KEY SAFE HANDLING
# ---------------------------------
openai.api_key = (
    st.secrets.get("OPENAI_API_KEY")  # Online (Streamlit Cloud)
    or os.getenv("OPENAI_API_KEY")     # Local environment variable
    or "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"  # fallback for local testing only
)

if not openai.api_key:
    st.error("❌ OpenAI API key not found. Please set OPENAI_API_KEY.")
    st.stop()

translator = GoogleTranslator(source="auto", target="km")

# ---------------------------------
# Extract text from PDF
# ---------------------------------
@st.cache_data(show_spinner=False)
def extract_text(file_bytes):
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text.strip()

# ---------------------------------
# OCR for scanned PDFs
# ---------------------------------
@st.cache_data(show_spinner=False)
def ocr_pdf(file_bytes):
    images = convert_from_bytes(file_bytes, dpi=OCR_DPI)
    result = ""
    for img in images:
        result += pytesseract.image_to_string(
            img, lang="eng+khm", config="--psm 6"
        ) + "\n"
    return result.strip()

# ---------------------------------
# Normalize text, preserve lines
# ---------------------------------
def normalize(text):
    lines = text.splitlines()
    return "\n".join(re.sub(r"\s{2,}", " ", l.rstrip()) for l in lines)

# ---------------------------------
# Fast translation: English → Khmer
# ---------------------------------
@st.cache_data(show_spinner=False)
def translate_fast(text):
    paragraphs = text.split("\n\n")
    output = []

    for para in paragraphs:
        if para.strip() == "":
            output.append("")
            continue
        try:
            if detect(para) == "en":
                chunks = [
                    para[i:i+TRANSLATE_LIMIT]
                    for i in range(0, len(para), TRANSLATE_LIMIT)
                ]
                translated = "".join(translator.translate(c) for c in chunks)
                output.append(translated)
            else:
                output.append(para)
        except:
            output.append(para)

    return "\n\n".join(output)

# ---------------------------------
# AI Khmer polishing
# ---------------------------------
def ai_polish_khmer(text):
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a professional Khmer language editor. "
                    "Improve clarity, grammar, and flow of Khmer text. "
                    "Only rewrite sentences if they sound unnatural or unclear. "
                    "Do NOT change meaning. Preserve paragraphs and line breaks."
                )
            },
            {
                "role": "user",
                "content": text
            }
        ],
        temperature=0.2
    )
    return response.choices[0].message.content.strip()

# ---------------------------------
# Export to DOCX
# ---------------------------------
def export_docx(text):
    doc = Document()
    doc.add_heading("English → Khmer (AI Polished)", level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------
# STREAMLIT UI
# ---------------------------------
st.set_page_config(page_title="AI Khmer PDF Translator", layout="centered")
st.title("🤖 PDF → Khmer → AI-Polished Word")

uploaded = st.file_uploader("📤 Upload PDF", type=["pdf"])

if uploaded:
    file_bytes = uploaded.read()

    with st.spinner("📄 Extracting text..."):
        text = extract_text(file_bytes)

    if len(text) < 30:
        with st.spinner("🖼️ Scanned PDF detected – OCR running..."):
            text = ocr_pdf(file_bytes)

    text = normalize(text)

    with st.spinner("🌐 Translating to Khmer..."):
        translated = translate_fast(text)

    with st.spinner("🤖 AI polishing Khmer language..."):
        polished = ai_polish_khmer(translated)

    st.subheader("📄 Final Khmer Output (AI-checked)")
    st.text_area("Result", polished, height=350)

    if st.button("⬇️ Export Word (.docx)"):
        docx_file = export_docx(polished)
        st.download_button(
            "📥 Download DOCX",
            docx_file,
            file_name="khmer_ai_polished.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("✅ Done – AI polished successfully")
